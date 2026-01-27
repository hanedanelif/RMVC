"""
Markdown dosyalarina detayli aciklamalar ekleyip Word'e cevirme
================================================================
Hoca icin her seyi anlasilir hale getiren script.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_enhanced_word_doc(md_file, docx_file, title):
    """Markdown'i detayli aciklamalarla Word'e cevir."""
    
    doc = Document()
    
    # Sayfa ayarlari
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Ana baslik
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aciklama kutusu
    p = doc.add_paragraph()
    p.add_run('SAYIN HOCAM,').bold = True
    doc.add_paragraph(
        'Bu rapor, matematiksel arka plan bilgisi olan ancak programlama deneyimi '
        'olmayan okuyucular icin hazirlanmistir. Her kavram detayli aciklanmis, '
        'her tablo ve grafik yorumlanmis, tum teknik terimler acilmistir.'
    )
    doc.add_paragraph()
    
    # Markdown dosyasini oku
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    print(f"[1/3] {os.path.basename(md_file)} okunuyor: {len(lines)} satir")
    
    i = 0
    in_code_block = False
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Kod blogu
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line if line else ' ')
            p.paragraph_format.left_indent = Inches(0.5)
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Basliklar
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Yatay cizgi
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            i += 1
            continue
        
        # Liste
        if line.startswith('- ') or (len(line) > 2 and line[0].isdigit() and line[1:3] == '. '):
            text = line[2:].strip() if line.startswith('- ') else line[line.index('.')+2:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet' if line.startswith('- ') else 'List Number')
            i += 1
            continue
        
        # Tablo
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # Ayirici satiri atla
            if all(set(cell.strip()) <= {'-', ' '} for cell in cells if cell.strip()):
                i += 1
                continue
            
            table_data.append(cells)
            i += 1
            
            # Sonraki satir tablo degilse tabloyu olustur
            if i >= len(lines) or '|' not in lines[i] or not lines[i].strip().startswith('|'):
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            cell.text = cell_text
                            
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    doc.add_paragraph()
                
                in_table = False
                table_data = []
            continue
        
        # Bos satir
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # Normal paragraf
        text = line.strip()
        parts = re.split(r'(\*\*.*?\*\*)', text)
        p = doc.add_paragraph()
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part:
                p.add_run(part)
        
        i += 1
    
    print(f"[2/3] Word belgesi olusturuldu")
    
    doc.save(docx_file)
    print(f"[3/3] Kaydedildi: {os.path.basename(docx_file)}")
    print(f"  - Boyut: {os.path.getsize(docx_file) / 1024:.2f} KB\n")

def main():
    """Ana fonksiyon - 3 dosyayi isle."""
    print("=" * 80)
    print("DETAYLI ACIKLAMALI WORD DOSYALARI OLUSTURULUYOR")
    print("=" * 80)
    print()
    
    base_dir = os.path.dirname(__file__)
    
    # Dosya listesi
    files = [
        {
            'md': 'FINAL_REPORT_FOR_PAPER.md',
            'docx': 'HOCA_ICIN_1_ANA_RAPOR.docx',
            'title': 'Iteratif RMVC Yakinsama Analizi - Ana Rapor'
        },
        {
            'md': 'UPDATED_CONVERGENCE_FINDINGS.md',
            'docx': 'HOCA_ICIN_2_SENTETIK_DETAY.docx',
            'title': 'Sentetik Veri Detayli Analiz'
        },
        {
            'md': 'BINARY_CONVERSION_ANALYSIS.md',
            'docx': 'HOCA_ICIN_3_BINARY_DONUSUM.docx',
            'title': 'Binary Donusum Yontemleri'
        }
    ]
    
    print(f"Toplam {len(files)} dosya islenecek:\n")
    
    for idx, file_info in enumerate(files, 1):
        print(f"[{idx}/{len(files)}] {file_info['md']} -> {file_info['docx']}")
        print("-" * 80)
        
        md_path = os.path.join(base_dir, file_info['md'])
        docx_path = os.path.join(base_dir, file_info['docx'])
        
        if not os.path.exists(md_path):
            print(f"  [UYARI] Dosya bulunamadi, atlanıyor...\n")
            continue
        
        create_enhanced_word_doc(md_path, docx_path, file_info['title'])
    
    print("=" * 80)
    print("TAMAMLANDI!")
    print("=" * 80)
    print()
    print("Hocaya gonderilecek Word dosyalari:")
    print("  1. HOCA_ICIN_1_ANA_RAPOR.docx")
    print("  2. HOCA_ICIN_2_SENTETIK_DETAY.docx")
    print("  3. HOCA_ICIN_3_BINARY_DONUSUM.docx")
    print()
    print("+ Ek dosyalar:")
    print("  - SUPPLEMENTARY_MATRICES_*.xlsx (tum matrisler)")
    print("  - SUPPLEMENTARY_CONFIGS_*.csv (konfigurasyonlar)")
    print("  - SUPPLEMENTARY_STATISTICS_*.csv (istatistikler)")
    print("  - analysis_plots/ (6 grafik)")

if __name__ == "__main__":
    main()
