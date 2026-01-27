"""
Markdown dosyasini Word formatina cevirme
==========================================
FINAL_REPORT_FOR_PAPER.md dosyasini .docx formatina cevirir.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_word(md_file, docx_file):
    """Markdown dosyasini Word'e cevir."""
    
    # Word belgesi olustur
    doc = Document()
    
    # Sayfa ayarlari
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Markdown dosyasini oku
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    print(f"[1/3] Markdown dosyasi okunuyor: {len(lines)} satir")
    
    i = 0
    in_code_block = False
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Kod blogu
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line if line else ' ')
            p.style = 'Normal'
            p.paragraph_format.left_indent = Inches(0.5)
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Baslik 1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Baslik 2 (## )
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            i += 1
            continue
        
        # Baslik 3 (### )
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Yatay cizgi
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            i += 1
            continue
        
        # Liste (- veya 1. )
        if line.startswith('- ') or (len(line) > 2 and line[0].isdigit() and line[1:3] == '. '):
            text = line[2:].strip() if line.startswith('- ') else line[line.index('.')+2:].strip()
            # Bold metni isle
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet' if line.startswith('- ') else 'List Number')
            i += 1
            continue
        
        # Tablo
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Tablo satirini isle
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # Ayirici satiri atla (|---|---|)
            if all(set(cell.strip()) <= {'-', ' '} for cell in cells if cell.strip()):
                i += 1
                continue
            
            table_data.append(cells)
            i += 1
            
            # Sonraki satir tablo degilse tabloyu olustur
            if i >= len(lines) or '|' not in lines[i] or not lines[i].strip().startswith('|'):
                if table_data:
                    # Tablo olustur
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            # Bold metni isle
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                            cell.text = cell_text
                            
                            # Baslik satiri (ilk satir) kalin yap
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    doc.add_paragraph()  # Bosluk
                
                in_table = False
                table_data = []
            continue
        
        # Bos satir
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # Normal paragraf
        text = line.strip()
        
        # Bold metni isle
        parts = re.split(r'(\*\*.*?\*\*)', text)
        p = doc.add_paragraph()
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part:
                p.add_run(part)
        
        i += 1
    
    print(f"[2/3] Word belgesi olusturuldu")
    
    # Kaydet
    doc.save(docx_file)
    print(f"[3/3] Dosya kaydedildi: {docx_file}")
    print(f"  - Dosya boyutu: {os.path.getsize(docx_file) / 1024:.2f} KB")

def main():
    """Ana fonksiyon."""
    print("=" * 80)
    print("MARKDOWN TO WORD CONVERTER")
    print("=" * 80)
    print()
    
    # Dosya yollari
    base_dir = os.path.dirname(__file__)
    md_file = os.path.join(base_dir, "FINAL_REPORT_FOR_PAPER.md")
    docx_file = os.path.join(base_dir, "FINAL_REPORT_FOR_PAPER.docx")
    
    if not os.path.exists(md_file):
        print(f"[HATA] Markdown dosyasi bulunamadi: {md_file}")
        return
    
    # Donustur
    parse_markdown_to_word(md_file, docx_file)
    
    print("\n" + "=" * 80)
    print("TAMAMLANDI!")
    print("=" * 80)
    print(f"\nWord dosyasi hazir: {os.path.basename(docx_file)}")
    print("Hocaya gonderebilirsiniz!")

if __name__ == "__main__":
    main()
